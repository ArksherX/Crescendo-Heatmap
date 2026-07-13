"""Regenerate WHITEPAPER.docx and WHITEPAPER.pdf from WHITEPAPER.md.

Usage:  python build_whitepaper_docs.py

Deps (already available in this environment): python-docx, markdown, weasyprint.
This is a clean render, not a match of any prior template — its job is to keep
the .docx/.pdf in sync with the Markdown source (which is the source of truth).
"""
from __future__ import annotations

import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent
MD = ROOT / "WHITEPAPER.md"
DOCX = ROOT / "WHITEPAPER.docx"
PDF = ROOT / "WHITEPAPER.pdf"

# ---------------------------------------------------------------- DOCX ----
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def _add_runs(paragraph, text: str) -> None:
    for seg in _INLINE.split(text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = paragraph.add_run(seg[2:-2]); r.bold = True
        elif seg.startswith("`") and seg.endswith("`"):
            r = paragraph.add_run(seg[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            paragraph.add_run(seg)


def build_docx(md: str) -> None:
    doc = Document()
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            i += 1
            continue

        # table block
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"\s*\|?[\s:-]+\|", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator
            body = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                body.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                _add_runs(table.rows[0].cells[j].paragraphs[0], h)
            for row in body:
                cells = table.add_row().cells
                for j in range(len(header)):
                    _add_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "")
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            if level == 1 and not doc.paragraphs:
                h = doc.add_heading("", level=0)
                _add_runs(h, m.group(2))
            else:
                h = doc.add_heading("", level=min(level, 4))
                _add_runs(h, m.group(2))
            i += 1
            continue

        # bullet / numbered list
        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^\s*[-*]\s+", "", line)); i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\s*\d+\.\s+", "", line)); i += 1
            continue

        if line.strip() == "" or set(line.strip()) == {"-"}:
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs(p, line)
        i += 1

    doc.save(str(DOCX))


# ----------------------------------------------------------------- PDF ----
import markdown as md_lib
from weasyprint import HTML

CSS = """
@page { size: A4; margin: 2.2cm; }
body { font-family: 'DejaVu Sans','Segoe UI',sans-serif; font-size: 10.5pt; line-height: 1.45; color:#1a1a1a; }
h1 { font-size: 20pt; border-bottom: 2px solid #333; padding-bottom: 4px; }
h2 { font-size: 15pt; margin-top: 20px; color:#111; }
h3 { font-size: 12.5pt; color:#222; }
code { font-family:'DejaVu Sans Mono',monospace; background:#f2f2f2; padding:1px 3px; font-size:9pt; }
pre { background:#f6f8fa; padding:10px; border-radius:6px; font-size:8.5pt; overflow-x:auto; }
table { border-collapse: collapse; width:100%; font-size:9.5pt; margin:10px 0; }
th,td { border:1px solid #bbb; padding:5px 8px; text-align:left; }
th { background:#eee; }
"""


def build_pdf(md: str) -> None:
    html_body = md_lib.markdown(md, extensions=["tables", "fenced_code"])
    html = f"<html><head><meta charset='utf-8'><style>{CSS}</style></head><body>{html_body}</body></html>"
    HTML(string=html).write_pdf(str(PDF))


TXT = ROOT / "WHITEPAPER.txt"


def build_txt(md: str) -> None:
    """Plain-text render (for upload systems that accept only .txt)."""
    out: list[str] = []
    lines = md.splitlines()
    i = 0
    _strip = lambda s: s.replace("**", "").replace("`", "")
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                out.append("    " + lines[i]); i += 1
            i += 1
            continue

        # table: keep rows, drop the |---| separator, clean pipes
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"\s*\|?[\s:-]+\|", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            out.append("    " + "  |  ".join(_strip(c) for c in header))
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                out.append("    " + "  |  ".join(_strip(c) for c in row)); i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level, txt = len(m.group(1)), _strip(m.group(2))
            if level == 1:
                out += ["", txt.upper(), "=" * len(txt)]
            elif level == 2:
                out += ["", txt, "-" * len(txt)]
            else:
                out += ["", txt]
            i += 1
            continue

        if line.strip() == "---":
            out.append("-" * 70); i += 1; continue

        out.append(_strip(line)); i += 1

    TXT.write_text("\n".join(out).strip() + "\n", encoding="utf-8")


def main() -> None:
    md = MD.read_text(encoding="utf-8")
    build_docx(md)
    print(f"wrote {DOCX.name}")
    build_pdf(md)
    print(f"wrote {PDF.name}")
    build_txt(md)
    print(f"wrote {TXT.name}")


if __name__ == "__main__":
    main()
